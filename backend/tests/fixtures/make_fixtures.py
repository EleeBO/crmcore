"""Script to generate binary test fixtures: sample.pdf, sample.xlsx, sample.docx."""

import io
import os

FIXTURES_DIR = os.path.dirname(__file__)


def make_pdf() -> bytes:
    """Create a minimal PDF with 2 pages for testing."""
    import fitz  # PyMuPDF

    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((50, 100), "Page 1: SLA Gold - RTO: 15 минут")
    page1.insert_text((50, 130), "Тариф Gold стоит $500 в месяц")
    page2 = doc.new_page()
    page2.insert_text((50, 100), "Page 2: Технические характеристики системы")
    page2.insert_text((50, 130), "Восстановление после сбоя: 15 минут по SLA Gold")
    return doc.tobytes()


def make_xlsx() -> bytes:
    """Create a minimal XLSX with pricing table."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Pricing"
    ws.append(["Plan", "Price", "RTO", "Features"])
    ws.append(["Gold", "$500", "15 min", "24/7 support"])
    ws.append(["Silver", "$200", "4 hours", "Business hours"])
    ws.append(["Bronze", "$50", "24 hours", "Email only"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_docx() -> bytes:
    """Create a minimal DOCX with section headers."""
    from docx import Document

    doc = Document()
    doc.add_heading("Переговорная стратегия", level=1)
    doc.add_paragraph("Основные тезисы для переговоров с клиентом.")
    doc.add_heading("Ценностное предложение", level=2)
    doc.add_paragraph("Наш продукт обеспечивает SLA Gold с RTO 15 минут.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    with open(os.path.join(FIXTURES_DIR, "sample.pdf"), "wb") as f:
        f.write(make_pdf())
    with open(os.path.join(FIXTURES_DIR, "sample.xlsx"), "wb") as f:
        f.write(make_xlsx())
    with open(os.path.join(FIXTURES_DIR, "sample.docx"), "wb") as f:
        f.write(make_docx())
    print("Fixtures generated.")
